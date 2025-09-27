import os
import uuid
import json
from pathlib import Path
from typing import Optional
import time

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile, status
import logging
import threading
from fastapi.responses import JSONResponse, PlainTextResponse, FileResponse

from ..config import settings
from ..auth import get_current_user
from ..tasks import process_video_task
from ..celery_app import celery_app

router = APIRouter()

def find_transcript_file(job_id: str) -> Optional[Path]:
    """Find transcript file by job_id, looking for both old UUID format and new filename format"""
    transcript_dir = Path(settings.transcript_dir)
    
    # First try the old UUID format
    old_format = transcript_dir / f"{job_id}.json"
    if old_format.exists():
        return old_format
    
    # Look for files that might match this job_id
    # This is a fallback - ideally we'd store the mapping
    for json_file in transcript_dir.glob("*.json"):
        if json_file.name != f"{job_id}.json":
            try:
                with open(json_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    if data.get("job_id") == job_id:
                        return json_file
            except:
                continue
    
    return None

def get_base_name_from_file(file_path: Path) -> str:
    """Extract base name from transcript file path"""
    return file_path.stem


logger = logging.getLogger(__name__)


def _enqueue_async(video_path: str, job_id: str, email: str | None):
    """Fire-and-forget enqueue with small retry loop to avoid failing HTTP."""
    import time
    attempts = 0
    last_err: Exception | None = None
    while attempts < 5:
        try:
            celery_app.send_task(
                name="process_video_task",
                args=[video_path, job_id, email],
                queue="celery",
            )
            logger.info("Enqueued job %s", job_id)
            return
        except Exception as e:  # noqa: BLE001
            last_err = e
            logger.warning("Enqueue attempt %s failed for %s: %s", attempts + 1, job_id, e)
            time.sleep(1 + attempts)
            attempts += 1
    if last_err:
        logger.error("Failed to enqueue job %s after retries: %s", job_id, last_err)


@router.post("/upload", status_code=202)
async def upload_video(file: UploadFile = File(...), user=Depends(get_current_user)):
    upload_dir = Path(settings.upload_dir)
    upload_dir.mkdir(parents=True, exist_ok=True)
    job_id = str(uuid.uuid4())
    video_path = upload_dir / f"{job_id}_{file.filename}"

    # Stream upload to disk to avoid loading entire file into memory
    try:
        with open(video_path, "wb") as out:
            while True:
                chunk = await file.read(1024 * 1024)  # 1 MiB
                if not chunk:
                    break
                out.write(chunk)
    finally:
        await file.close()

    # Enqueue in a background thread with retries; always return 202
    # persist a simple task mapping file so /status can find the task later if needed
    task_map_dir = Path(settings.transcript_dir) / "_tasks"
    task_map_dir.mkdir(parents=True, exist_ok=True)
    (task_map_dir / f"{job_id}.task").write_text("pending", encoding="utf-8")

    threading.Thread(target=_enqueue_async, args=(str(video_path), job_id, user.get("email")), daemon=True).start()
    return {"job_id": job_id}


@router.get("/status/{job_id}")
async def job_status(job_id: str, user=Depends(get_current_user)):
    transcript_json = find_transcript_file(job_id)
    if transcript_json and transcript_json.exists():
        return {"job_id": job_id, "status": "completed", "percent": 100, "stage": "done"}
    
    # Check if job is still processing by looking at task file
    task_map = Path(settings.transcript_dir) / "_tasks" / f"{job_id}.task"
    if not task_map.exists():
        # No task file means job doesn't exist or failed
        return {"job_id": job_id, "status": "not_found", "percent": 0, "stage": "unknown"}
    
    stage = "transcribe"
    percent = 10
    try:
        raw = task_map.read_text(encoding="utf-8").strip()
        parts = raw.split("|", 2)
        if len(parts) >= 2:
            stage = parts[0]
            percent = int(parts[1])
    except Exception:
        pass
    
    # Check if job has been stuck for too long (more than 30 minutes)
    try:
        file_age = time.time() - task_map.stat().st_mtime
        if file_age > 1800:  # 30 minutes
            return {"job_id": job_id, "status": "timeout", "percent": percent, "stage": stage}
    except Exception:
        pass
    
    return {"job_id": job_id, "status": "processing", "percent": percent, "stage": stage}

@router.get("/logs/{job_id}")
async def job_logs(job_id: str, user=Depends(get_current_user)):
    log_path = Path(settings.transcript_dir) / "_tasks" / f"{job_id}.log"
    if not log_path.exists():
        return PlainTextResponse("", status_code=200)
    return PlainTextResponse(log_path.read_text(encoding="utf-8"))


@router.get("/transcript/{job_id}")
async def get_transcript(job_id: str, format: Optional[str] = "json", user=Depends(get_current_user)):
	transcript_dir = Path(settings.transcript_dir)
	json_path = find_transcript_file(job_id)
	if not json_path:
		raise HTTPException(status_code=404, detail="Transcript not found")
	
	base_name = get_base_name_from_file(json_path)
	txt_path = transcript_dir / f"{base_name}.txt"
	sum_path = transcript_dir / f"{base_name}.summary.txt"
	if format == "json":
		if not json_path.exists():
			raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Not found")
		return JSONResponse(content=json.loads(json_path.read_text(encoding="utf-8")))
	elif format == "txt":
		if not txt_path.exists():
			raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Not found")
		return PlainTextResponse(content=txt_path.read_text(encoding="utf-8"))
	elif format == "summary":
		if not sum_path.exists():
			raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Not found")
		return PlainTextResponse(content=sum_path.read_text(encoding="utf-8"))
	else:
		raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unsupported format")


@router.get("/export/{job_id}")
async def export_transcript(job_id: str, format: str = "pdf", user=Depends(get_current_user)):
	"""Export transcript in various formats: pdf, docx"""
	transcript_dir = Path(settings.transcript_dir)
	json_path = transcript_dir / f"{job_id}.json"
	
	if not json_path.exists():
		raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Transcript not found")
	
	data = json.loads(json_path.read_text(encoding="utf-8"))
	
	if format == "pdf":
		return _export_pdf(data, job_id)
	elif format == "docx":
		return _export_docx(data, job_id)
	else:
		raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unsupported export format")


def _export_pdf(data, job_id):
	"""Export transcript as PDF"""
	try:
		from reportlab.lib.pagesizes import letter, A4
		from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
		from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
		from reportlab.lib.units import inch
		from reportlab.lib import colors
		from io import BytesIO
		
		buffer = BytesIO()
		doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
		story = []
		styles = getSampleStyleSheet()
		
		# Title
		title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=16, spaceAfter=30)
		story.append(Paragraph(f"Meeting Transcript - {job_id}", title_style))
		story.append(Spacer(1, 12))
		
		# Summary if available
		if 'summary' in data and data['summary']:
			story.append(Paragraph("Summary", styles['Heading2']))
			story.append(Paragraph(data['summary'].replace('\n', '<br/>'), styles['Normal']))
			story.append(Spacer(1, 12))
		
		# Transcript segments
		story.append(Paragraph("Transcript", styles['Heading2']))
		story.append(Spacer(1, 12))
		
		for segment in data.get('segments', []):
			start_time = segment.get('start', 0)
			end_time = segment.get('end', 0)
			text = segment.get('text', '')
			speaker = segment.get('speaker', 'Unknown')
			
			time_str = f"[{start_time:.1f}s - {end_time:.1f}s]"
			speaker_str = f"<b>{speaker}:</b>" if speaker != 'Unknown' else ""
			
			story.append(Paragraph(f"{time_str} {speaker_str} {text}", styles['Normal']))
			story.append(Spacer(1, 6))
		
		doc.build(story)
		buffer.seek(0)
		
		return FileResponse(
			path=buffer,
			filename=f"transcript_{job_id}.pdf",
			media_type="application/pdf",
			background=None
		)
	except ImportError:
		raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="PDF export not available - missing reportlab")


def _export_docx(data, job_id):
	"""Export transcript as DOCX"""
	try:
		from docx import Document
		from docx.shared import Inches
		from io import BytesIO
		
		doc = Document()
		doc.add_heading(f'Meeting Transcript - {job_id}', 0)
		
		# Summary
		if 'summary' in data and data['summary']:
			doc.add_heading('Summary', level=1)
			doc.add_paragraph(data['summary'])
		
		# Transcript
		doc.add_heading('Transcript', level=1)
		
		for segment in data.get('segments', []):
			start_time = segment.get('start', 0)
			end_time = segment.get('end', 0)
			text = segment.get('text', '')
			speaker = segment.get('speaker', 'Unknown')
			
			time_str = f"[{start_time:.1f}s - {end_time:.1f}s]"
			speaker_str = f"{speaker}:" if speaker != 'Unknown' else ""
			
			doc.add_paragraph(f"{time_str} {speaker_str} {text}")
		
		buffer = BytesIO()
		doc.save(buffer)
		buffer.seek(0)
		
		return FileResponse(
			path=buffer,
			filename=f"transcript_{job_id}.docx",
			media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
			background=None
		)
	except ImportError:
		raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="DOCX export not available - missing python-docx")




@router.put("/transcript/{job_id}/speakers")
async def update_speakers(job_id: str, speakers: dict, user=Depends(get_current_user)):
	"""Update speaker names for a transcript"""
	transcript_dir = Path(settings.transcript_dir)
	json_path = transcript_dir / f"{job_id}.json"
	
	if not json_path.exists():
		raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Transcript not found")
	
	# Load current transcript
	data = json.loads(json_path.read_text(encoding="utf-8"))
	
	# Update speaker names in segments
	for segment in data.get('segments', []):
		old_speaker = segment.get('speaker', '')
		if old_speaker in speakers:
			segment['speaker'] = speakers[old_speaker]
	
	# Save updated transcript
	json_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
	
	# Update TXT file
	txt_path = transcript_dir / f"{job_id}.txt"
	if txt_path.exists():
		lines = []
		for seg in data["segments"]:
			spk = seg.get("speaker", "")
			text = seg.get("text", "")
			lines.append((f"[{spk}] " if spk else "") + text)
		txt_path.write_text("\n".join(lines), encoding="utf-8")
	
	return {"status": "success", "message": "Speaker names updated"}


@router.get("/search")
async def search_transcripts(query: str, user=Depends(get_current_user)):
	"""Search across all transcripts"""
	transcript_dir = Path(settings.transcript_dir)
	results = []
	
	if not query or not query.strip():
		return {"query": query, "results": [], "total": 0}
	
	query_lower = query.lower().strip()
	
	# Search through all JSON files
	for json_file in transcript_dir.glob("*.json"):
		try:
			data = json.loads(json_file.read_text(encoding="utf-8"))
			job_id = json_file.stem
			
			# Search in segments
			matching_segments = []
			for segment in data.get('segments', []):
				text = segment.get('text', '').lower()
				if query_lower in text:
					matching_segments.append(segment)
			
			# Search in summary
			summary_matches = []
			if 'summary' in data and data['summary']:
				if query_lower in data['summary'].lower():
					summary_matches.append(data['summary'])
			
			if matching_segments or summary_matches:
				# Try to get original filename from upload directory
				upload_dir = Path(settings.upload_dir)
				filename = None
				for upload_file in upload_dir.glob(f"{job_id}_*"):
					if upload_file.is_file():
						filename = upload_file.name.replace(f"{job_id}_", "")
						break
				
				results.append({
					"job_id": job_id,
					"filename": filename,
					"segments": matching_segments,
					"summary_matches": summary_matches,
					"total_segments": len(data.get('segments', [])),
					"language": data.get('language', 'unknown')
				})
		except Exception as e:
			print(f"Error processing {json_file}: {e}")
			continue
	
	return {"query": query, "results": results, "total": len(results)}


@router.get("/jobs")
async def get_jobs(user=Depends(get_current_user)):
	"""Get list of all jobs with their status and filenames"""
	transcript_dir = Path(settings.transcript_dir)
	jobs = []
	
	# Get all JSON files (completed jobs)
	for json_file in transcript_dir.glob("*.json"):
		try:
			data = json.loads(json_file.read_text(encoding="utf-8"))
			job_id = json_file.stem
			
			# Try to get original filename from upload directory
			upload_dir = Path(settings.upload_dir)
			filename = None
			for upload_file in upload_dir.glob(f"{job_id}_*"):
				if upload_file.is_file():
					filename = upload_file.name.replace(f"{job_id}_", "")
					break
			
			jobs.append({
				"job_id": job_id,
				"filename": filename,
				"status": "completed",
				"language": data.get('language', 'unknown'),
				"created_at": json_file.stat().st_mtime
			})
		except Exception:
			continue
	
	# Get processing jobs from task files
	task_dir = Path(settings.transcript_dir) / "_tasks"
	if task_dir.exists():
		for task_file in task_dir.glob("*.task"):
			job_id = task_file.stem
			try:
				# Skip if already completed
				if any(job['job_id'] == job_id for job in jobs):
					continue
				
				# Try to get filename from upload directory
				upload_dir = Path(settings.upload_dir)
				filename = None
				for upload_file in upload_dir.glob(f"{job_id}_*"):
					if upload_file.is_file():
						filename = upload_file.name.replace(f"{job_id}_", "")
						break
				
				jobs.append({
					"job_id": job_id,
					"filename": filename,
					"status": "processing",
					"language": "unknown",
					"created_at": task_file.stat().st_mtime
				})
			except Exception:
				continue
	
	# Sort by creation time (newest first)
	jobs.sort(key=lambda x: x['created_at'], reverse=True)
	
	return jobs


@router.get("/analytics")
async def get_analytics(user=Depends(get_current_user)):
	"""Get usage statistics and performance metrics"""
	transcript_dir = Path(settings.transcript_dir)
	
	# Count total transcripts
	total_transcripts = len(list(transcript_dir.glob("*.json")))
	
	# Calculate total duration
	total_duration = 0
	language_stats = {}
	file_sizes = []
	
	for json_file in transcript_dir.glob("*.json"):
		try:
			data = json.loads(json_file.read_text(encoding="utf-8"))
			
			# Calculate duration
			segments = data.get('segments', [])
			if segments:
				last_segment = max(segments, key=lambda x: x.get('end', 0))
				total_duration += last_segment.get('end', 0)
			
			# Language stats
			lang = data.get('language', 'unknown')
			language_stats[lang] = language_stats.get(lang, 0) + 1
			
			# File size
			file_sizes.append(json_file.stat().st_size)
		except Exception:
			continue
	
	# Calculate averages
	avg_duration = total_duration / total_transcripts if total_transcripts > 0 else 0
	avg_file_size = sum(file_sizes) / len(file_sizes) if file_sizes else 0
	
	return {
		"total_transcripts": total_transcripts,
		"total_duration_hours": total_duration / 3600,
		"average_duration_minutes": avg_duration / 60,
		"language_distribution": language_stats,
		"average_file_size_mb": avg_file_size / (1024 * 1024),
		"storage_used_mb": sum(file_sizes) / (1024 * 1024)
	}

@router.delete("/transcripts/clear")
async def clear_all_transcripts(user=Depends(get_current_user)):
	"""Clear all user's transcripts with confirmation"""
	transcript_dir = Path(settings.transcript_dir)
	
	# Count files to be deleted
	json_files = list(transcript_dir.glob("*.json"))
	txt_files = list(transcript_dir.glob("*.txt"))
	summary_files = list(transcript_dir.glob("*.summary.txt"))
	pdf_files = list(transcript_dir.glob("*.pdf"))
	docx_files = list(transcript_dir.glob("*.docx"))
	
	total_files = len(json_files) + len(txt_files) + len(summary_files) + len(pdf_files) + len(docx_files)
	
	if total_files == 0:
		return {"message": "No transcripts found to delete", "deleted_count": 0}
	
	# Delete all transcript files
	deleted_count = 0
	for file_list in [json_files, txt_files, summary_files, pdf_files, docx_files]:
		for file_path in file_list:
			try:
				file_path.unlink()
				deleted_count += 1
			except Exception as e:
				logger.error(f"Error deleting file {file_path}: {e}")
	
	# Also clear task files
	task_dir = transcript_dir / "_tasks"
	if task_dir.exists():
		for task_file in task_dir.glob("*.task"):
			try:
				task_file.unlink()
			except Exception as e:
				logger.error(f"Error deleting task file {task_file}: {e}")
		
		for log_file in task_dir.glob("*.log"):
			try:
				log_file.unlink()
			except Exception as e:
				logger.error(f"Error deleting log file {log_file}: {e}")
	
	return {
		"message": f"Successfully deleted {deleted_count} transcript files",
		"deleted_count": deleted_count
	}
